"""
AI GOVERNANCE & POLICY ENGINEERING PLATFORM
Enterprise-Grade Agentic AI Application
Single file: app.py
Full Demo Mode support (no API key required)
Production-ready, modular, multi-agent architecture.
"""

import streamlit as st
import openai
import os
import json
import time
import datetime
from typing import List, Dict, Any, Optional, Tuple, Union
from io import BytesIO
import base64
import re
import textwrap
import logging
from copy import deepcopy

# Document parsing
import docx as docx_module
from PyPDF2 import PdfReader
import markdown as md_lib

# Document generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from fpdf import FPDF

# For diff viewer
import difflib
import pandas as pd

# --------------------------------------------------------------------------------
# Logging Configuration
# --------------------------------------------------------------------------------
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# --------------------------------------------------------------------------------
# Configuration & Constants
# --------------------------------------------------------------------------------
class Config:
    # Model settings (used only when API key is provided)
    MODEL_NAME = "gpt-4o"
    TEMPERATURE = 0.2
    MAX_TOKENS = 4096
    REQUEST_TIMEOUT = 120

    # Application metadata
    APP_NAME = "PolicyAI Enterprise"
    APP_VERSION = "2.0.0"
    SUPPORTED_INPUT_FORMATS = ["pdf", "docx", "txt", "md"]
    OUTPUT_FORMATS = ["docx", "pdf", "md", "json"]

    # Compliance frameworks
    FRAMEWORKS = [
        "ISO 27001",
        "COBIT 2019",
        "NIST SP 800-53",
        "CIS Controls v8",
        "GDPR",
        "UU PDP Indonesia",
        "PCI DSS",
        "HIPAA",
        "SOX",
        "Custom"
    ]

    # Industry types
    INDUSTRIES = [
        "Financial Services",
        "Healthcare",
        "Technology",
        "Government",
        "Manufacturing",
        "Energy",
        "Education",
        "Retail",
        "Telecommunications",
        "Other"
    ]

    # Maturity levels
    MATURITY_LEVELS = ["Initial", "Repeatable", "Defined", "Managed", "Optimizing"]

    # Organization scale
    ORG_SCALES = ["Small (<100 employees)", "Medium (100-1000)", "Large (1000-10000)", "Enterprise (10000+)"]

# --------------------------------------------------------------------------------
# Mock Data Generator for Demo Mode
# --------------------------------------------------------------------------------
def generate_mock_agent_outputs(state: PolicyState) -> Dict[str, Any]:
    """Return realistic dummy outputs for all agents, using user inputs for flavor."""
    org = state.user_inputs.get("org_type", "Organization")
    policy_name = state.user_inputs.get("policy_name", "Policy")
    industry = state.user_inputs.get("industry", "Technology")
    maturity = state.user_inputs.get("cyber_maturity", "Defined")

    return {
        "Requirement Analyst": {
            "policy_type": "Information Security Policy",
            "scope": f"All {org} employees, contractors, and third-party users with access to information assets.",
            "objectives": [
                f"Protect the confidentiality, integrity, and availability of {org} information assets.",
                "Ensure compliance with ISO 27001 and applicable regulations.",
                "Establish a risk-based approach to information security management."
            ],
            "context_analysis": f"As a {industry} organization, {org} faces threats related to data breaches, ransomware, and insider threats.",
            "inferred_requirements": [
                "Access control policy",
                "Incident response plan",
                "Data classification standard",
                "Acceptable use policy"
            ],
            "suggested_controls_focus": ["Access Management", "Encryption", "Monitoring"]
        },
        "Governance & Compliance": {
            "compliance_mapping": [
                {"framework": "ISO 27001", "clause": "A.9.1.1", "control_id": "Access control policy", "status": "aligned"},
                {"framework": "ISO 27001", "clause": "A.12.1.1", "control_id": "Incident management", "status": "gap"},
                {"framework": "NIST SP 800-53", "clause": "AC-1", "control_id": "Access control policy and procedures", "status": "aligned"}
            ],
            "governance_recommendations": [
                "Establish an Information Security Steering Committee.",
                "Define clear roles for Data Owner, Custodian, and User.",
                "Implement a quarterly policy review cycle."
            ],
            "regulatory_gaps": ["Lack of incident reporting procedure aligned with UU PDP Indonesia."],
            "required_evidence": ["Approved access control matrix", "Incident response test records", "Policy acknowledgment receipts"]
        },
        "Risk Assessment": {
            "risks": [
                {
                    "id": "R1",
                    "description": "Unauthorized access to sensitive data",
                    "category": "Confidentiality",
                    "likelihood": "medium",
                    "impact": "high",
                    "criticality": "high",
                    "suggested_controls": ["Multi-factor authentication", "Role-based access control"]
                },
                {
                    "id": "R2",
                    "description": "Ransomware attack via phishing",
                    "category": "Malware",
                    "likelihood": "high",
                    "impact": "high",
                    "criticality": "critical",
                    "suggested_controls": ["Advanced email filtering", "Security awareness training", "Endpoint detection and response"]
                }
            ],
            "risk_appetite_statement": f"{org} maintains a low risk appetite for information security, accepting only minor operational risks.",
            "control_prioritization": ["MFA", "EDR", "Employee training"]
        },
        "Technical Control": {
            "access_control": ["MFA for all remote access", "RBAC with quarterly reviews"],
            "encryption_standards": "AES-256 for data at rest, TLS 1.3 for data in transit",
            "network_security": ["Next-gen firewall with IPS", "Network segmentation", "Zero Trust architecture"],
            "endpoint_protection": ["EDR solution", "Disk encryption", "Application whitelisting"],
            "monitoring_and_logging": ["SIEM with real-time alerting", "Log retention 1 year"],
            "incident_response": "Formal IR plan with roles and playbooks, tested annually",
            "vulnerability_management": "Monthly vulnerability scans, critical patching within 7 days",
            "security_baseline": "CIS Benchmark Level 1 for all servers"
        },
        "Legal & Policy Writer": {
            "purpose": f"The purpose of this {policy_name} is to establish the framework for protecting {org}'s information assets...",
            "scope": f"This policy applies to all employees, contractors, and third parties...",
            "definitions": {
                "Asset": "Anything that has value to the organization.",
                "Confidentiality": "Ensuring that information is accessible only to those authorized."
            },
            "governance_structure": "The Chief Information Security Officer (CISO) is responsible for...",
            "roles_and_responsibilities": "All employees must adhere to this policy and report any incidents immediately.",
            "policy_statements": [
                "Access to information must be based on business need-to-know.",
                "All devices must be encrypted.",
                "Regular security awareness training is mandatory."
            ],
            "operational_controls": "Change management, backup and recovery procedures...",
            "monitoring_mechanism": "Continuous monitoring via SIEM, monthly KPI reports.",
            "exception_handling": "Exceptions require formal approval from CISO.",
            "violation_sanctions": "Violations may result in disciplinary action up to termination.",
            "kpi_kri": ["Number of security incidents per month", "Patch compliance percentage", "Phishing test success rate"]
        },
        "Audit Readiness": {
            "evidence_requirements": [
                {"control_area": "Access Control", "evidence_type": "User access review reports", "frequency": "Quarterly"},
                {"control_area": "Encryption", "evidence_type": "Encryption standard documentation", "frequency": "Annual"}
            ],
            "kpi_kri_recommendations": [
                {"metric": "Incident resolution time", "target": "< 4 hours for P1", "review_frequency": "Monthly"},
                {"metric": "Patch compliance", "target": "> 98%", "review_frequency": "Weekly"}
            ],
            "monitoring_mechanism": "Automated dashboards with alerts for deviations.",
            "audit_trail_design": "Centralized logging with integrity protection.",
            "review_cycle": "Policy review every 12 months or after major incident."
        },
        "Reviewer": {
            "conflicts": [],
            "ambiguities": ["Scope could be clarified for remote workers."],
            "overlaps": ["Incident response overlaps with business continuity plan."],
            "terminology_issues": ["Consistent use of 'must' vs 'shall'."],
            "feasibility_concerns": ["Implementation of zero trust may require significant investment."],
            "overall_quality_score": 0.85
        },
        "Legacy Policy Analysis": {
            "structure_summary": "Old policy with 8 sections, missing technical controls.",
            "sections": ["Purpose", "Scope", "Policy", "Enforcement"],
            "obsolete_clauses": ["Use of floppy disks", "Password rotation every 30 days without MFA"],
            "weaknesses": ["No incident response plan", "Vague access control statements"],
            "duplicates": ["Two sections on password policy"],
            "governance_maturity_estimate": "Repeatable",
            "critical_issues": ["Missing encryption requirements"]
        },
        "Gap Analysis": {
            "gap_matrix": [
                {
                    "existing_clause": "Password policy",
                    "identified_gap": "No MFA requirement",
                    "risk_level": "High",
                    "recommendation": "Mandate MFA for all remote access",
                    "target_framework": "ISO 27001 A.9.4.2",
                    "priority": "1"
                },
                {
                    "existing_clause": "Data backup",
                    "identified_gap": "No off-site backup requirement",
                    "risk_level": "Medium",
                    "recommendation": "Define off-site backup with 30-day retention",
                    "target_framework": "CIS Controls v8 11.2",
                    "priority": "2"
                }
            ],
            "summary": "Two high-risk and one medium-risk gaps identified."
        },
        "Policy Modernization": {
            "modernized_policy": {
                "purpose": f"The purpose of this {policy_name} is to establish the framework for protecting {org}'s information assets...",
                "scope": f"This policy applies to all employees, contractors, and third parties...",
                "definitions": {
                    "Asset": "Anything that has value to the organization.",
                    "Confidentiality": "Ensuring that information is accessible only to those authorized."
                },
                "governance_structure": "The Chief Information Security Officer (CISO) is responsible for...",
                "roles_and_responsibilities": "All employees must adhere to this policy and report any incidents immediately.",
                "policy_statements": [
                    "Access to information must be based on business need-to-know.",
                    "All devices must be encrypted.",
                    "Regular security awareness training is mandatory."
                ],
                "operational_controls": "Change management, backup and recovery procedures...",
                "monitoring_mechanism": "Continuous monitoring via SIEM, monthly KPI reports.",
                "exception_handling": "Exceptions require formal approval from CISO.",
                "violation_sanctions": "Violations may result in disciplinary action up to termination.",
                "kpi_kri": ["Number of security incidents per month", "Patch compliance percentage", "Phishing test success rate"]
            },
            "change_log": [
                {
                    "old_clause": "Password must be changed every 30 days.",
                    "new_clause": "Passwords must be complex; MFA is mandatory.",
                    "reason": "NIST 800-63B recommends against periodic changes without cause.",
                    "improvement_category": "Security Control",
                    "compliance_impact": "High"
                },
                {
                    "old_clause": "Backups are performed weekly.",
                    "new_clause": "Backups performed daily, with off-site copy and monthly restore test.",
                    "reason": "Improved RPO/RTO and resilience.",
                    "improvement_category": "Operational Control",
                    "compliance_impact": "Medium"
                }
            ],
            "modernization_report": {
                "executive_summary": "The policy has been significantly strengthened.",
                "major_improvements": ["Added MFA", "Incident response plan", "Data classification"],
                "governance_improvement": "Defined CISO role and committee.",
                "compliance_improvement": "Aligned with ISO 27001:2022.",
                "security_improvement": "Zero Trust principles incorporated.",
                "audit_readiness_improvement": "Clear evidence requirements and KPIs.",
                "next_actions": ["Conduct awareness training", "Implement EDR", "Quarterly audit"]
            }
        }
    }

# --------------------------------------------------------------------------------
# AI Client Setup
# --------------------------------------------------------------------------------
def get_openai_client() -> Optional[openai.OpenAI]:
    """Returns OpenAI client using session state API key or env variable."""
    api_key = st.session_state.get("openai_api_key", "") or os.getenv("OPENAI_API_KEY")
    if not api_key:
        return None
    try:
        client = openai.OpenAI(api_key=api_key)
        return client
    except Exception as e:
        logger.error(f"Failed to initialize OpenAI client: {e}")
        return None

# --------------------------------------------------------------------------------
# Document Parsing Utilities
# --------------------------------------------------------------------------------
class DocumentParser:
    """Intelligent parser for PDF, DOCX, TXT, MD with structure preservation."""

    @staticmethod
    def parse_pdf(file_bytes: bytes) -> str:
        try:
            reader = PdfReader(BytesIO(file_bytes))
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
            return text
        except Exception as e:
            logger.error(f"PDF parsing error: {e}")
            raise ValueError(f"Failed to parse PDF: {e}")

    @staticmethod
    def parse_docx(file_bytes: bytes) -> str:
        try:
            doc = docx_module.Document(BytesIO(file_bytes))
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    full_text.append(" | ".join(row_text))
            return "\n".join(full_text)
        except Exception as e:
            logger.error(f"DOCX parsing error: {e}")
            raise ValueError(f"Failed to parse DOCX: {e}")

    @staticmethod
    def parse_markdown(file_bytes: bytes) -> str:
        try:
            md_content = file_bytes.decode('utf-8')
            html = md_lib.markdown(md_content)
            clean = re.sub(r'<[^>]+>', '', html)
            return clean
        except Exception as e:
            logger.error(f"Markdown parsing error: {e}")
            raise ValueError(f"Failed to parse Markdown: {e}")

    @staticmethod
    def parse_txt(file_bytes: bytes) -> str:
        try:
            return file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            return file_bytes.decode('latin-1')
        except Exception as e:
            logger.error(f"TXT parsing error: {e}")
            raise ValueError(f"Failed to parse text file: {e}")

    @classmethod
    def parse_document(cls, file_bytes: bytes, file_type: str) -> str:
        file_type = file_type.lower()
        if file_type == "pdf":
            return cls.parse_pdf(file_bytes)
        elif file_type == "docx":
            return cls.parse_docx(file_bytes)
        elif file_type in ["txt", "text"]:
            return cls.parse_txt(file_bytes)
        elif file_type in ["md", "markdown"]:
            return cls.parse_markdown(file_bytes)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

# --------------------------------------------------------------------------------
# Semantic Chunking
# --------------------------------------------------------------------------------
class SemanticChunker:
    CHUNK_SIZE = 1500
    CHUNK_OVERLAP = 200

    @classmethod
    def chunk_by_headings(cls, text: str) -> List[Dict[str, str]]:
        lines = text.split('\n')
        chunks = []
        current_chunk = ""
        current_heading = "Introduction"
        for line in lines:
            if (line.isupper() and len(line) > 5) or (line.strip().endswith(':') and len(line.strip()) < 60):
                if current_chunk.strip():
                    chunks.append({"heading": current_heading, "content": current_chunk.strip()})
                current_heading = line.strip().rstrip(':')
                current_chunk = ""
            else:
                current_chunk += line + "\n"
        if current_chunk.strip():
            chunks.append({"heading": current_heading, "content": current_chunk.strip()})
        return chunks

    @classmethod
    def chunk_semantic(cls, text: str) -> List[str]:
        chunks = []
        start = 0
        text_len = len(text)
        while start < text_len:
            end = start + cls.CHUNK_SIZE
            chunks.append(text[start:end])
            start = end - cls.CHUNK_OVERLAP
        return chunks

# --------------------------------------------------------------------------------
# Policy State Management
# --------------------------------------------------------------------------------
class PolicyState:
    def __init__(self):
        self.mode = "generate"
        self.user_inputs = {}
        self.raw_document = ""
        self.parsed_structure = {}
        self.chunks = []
        self.agent_outputs = {}
        self.metadata = {
            "start_time": time.time(),
            "execution_steps": [],
            "errors": []
        }
        self.final_policy_json = {}
        self.scores = {
            "compliance": 0.0,
            "modernization": 0.0,
            "governance_maturity": 0.0,
            "policy_quality": 0.0
        }
        self.gap_matrix = []
        self.change_log = []
        self.modernization_report = {}

# --------------------------------------------------------------------------------
# Base Agent Class
# --------------------------------------------------------------------------------
class BaseAgent:
    def __init__(self, name: str, system_prompt: str, client: Optional[openai.OpenAI]):
        self.name = name
        self.system_prompt = system_prompt
        self.client = client
        self.model = Config.MODEL_NAME

    def _call_openai(self, messages: List[Dict], response_format: Optional[Dict] = None) -> str:
        if not self.client:
            raise RuntimeError("OpenAI client not available.")
        try:
            kwargs = {
                "model": self.model,
                "messages": messages,
                "temperature": Config.TEMPERATURE,
                "max_tokens": Config.MAX_TOKENS,
                "timeout": Config.REQUEST_TIMEOUT,
            }
            if response_format:
                kwargs["response_format"] = response_format
            response = self.client.chat.completions.create(**kwargs)
            return response.choices[0].message.content
        except Exception as e:
            logger.error(f"Agent {self.name} API call failed: {e}")
            raise RuntimeError(f"Agent {self.name}: {str(e)}")

    def run(self, state: PolicyState, additional_context: str = "") -> Dict[str, Any]:
        raise NotImplementedError

    def _build_messages(self, user_message: str) -> List[Dict]:
        return [
            {"role": "system", "content": self.system_prompt},
            {"role": "user", "content": user_message}
        ]

    def _parse_json_response(self, raw_response: str) -> Dict:
        try:
            return json.loads(raw_response)
        except json.JSONDecodeError:
            json_match = re.search(r'```(?:json)?\s*([\s\S]*?)\s*```', raw_response)
            if json_match:
                return json.loads(json_match.group(1))
            else:
                raise ValueError(f"Agent {self.name} did not return valid JSON.")

# --------------------------------------------------------------------------------
# Specific Agent Implementations (identical to previous, but constructor receives client)
# --------------------------------------------------------------------------------
class RequirementAnalystAgent(BaseAgent):
    def __init__(self, client):
        super().__init__(
            name="Requirement Analyst",
            system_prompt=textwrap.dedent("""\
                You are a senior enterprise governance analyst. Based on user inputs, analyze and define:
                - policy type
                - policy scope
                - key objectives
                - organizational context implications
                - inferred missing context
                Output structured JSON with fields: policy_type, scope, objectives, context_analysis, inferred_requirements, suggested_controls_focus
            """),
            client=client
        )

    def run(self, state: PolicyState, additional_context: str = "") -> Dict[str, Any]:
        user_msg = f"User Inputs: {json.dumps(state.user_inputs, indent=2)}\n\nAdditional Context: {additional_context}"
        messages = self._build_messages(user_msg)
        raw = self._call_openai(messages, response_format={"type": "json_object"})
        return self._parse_json_response(raw)

class GovernanceComplianceAgent(BaseAgent):
    def __init__(self, client):
        super().__init__(
            name="Governance & Compliance",
            system_prompt=textwrap.dedent("""\
                You are a GRC expert. Validate alignment against selected frameworks.
                Provide compliance mapping, governance recommendations, regulatory gaps, required evidence.
                Output JSON.
            """),
            client=client
        )

    def run(self, state: PolicyState, additional_context: str = "") -> Dict[str, Any]:
        req_output = state.agent_outputs.get("Requirement Analyst", {})
        user_msg = f"Requirements: {json.dumps(req_output)}\nFrameworks selected: {state.user_inputs.get('main_regulation','')}, {state.user_inputs.get('existing_framework','')}\nAdditional: {additional_context}"
        messages = self._build_messages(user_msg)
        raw = self._call_openai(messages, response_format={"type": "json_object"})
        return self._parse_json_response(raw)

class RiskAssessmentAgent(BaseAgent):
    def __init__(self, client):
        super().__init__(
            name="Risk Assessment",
            system_prompt=textwrap.dedent("""\
                You are a risk management specialist. Perform risk identification and classification.
                Output JSON with risks, risk_appetite_statement, control_prioritization.
            """),
            client=client
        )

    def run(self, state: PolicyState, additional_context: str = "") -> Dict[str, Any]:
        context = json.dumps(state.agent_outputs.get("Requirement Analyst", {}))
        user_msg = f"Policy context: {context}\nIndustry: {state.user_inputs.get('industry','')}\nAdditional: {additional_context}"
        messages = self._build_messages(user_msg)
        raw = self._call_openai(messages, response_format={"type": "json_object"})
        return self._parse_json_response(raw)

class TechnicalControlAgent(BaseAgent):
    def __init__(self, client):
        super().__init__(
            name="Technical Control",
            system_prompt=textwrap.dedent("""\
                You are a cybersecurity architect. Recommend realistic technical controls aligned with maturity.
                Output JSON with access_control, encryption_standards, network_security, endpoint_protection, monitoring_and_logging, incident_response, vulnerability_management, security_baseline.
            """),
            client=client
        )

    def run(self, state: PolicyState, additional_context: str = "") -> Dict[str, Any]:
        risks = state.agent_outputs.get("Risk Assessment", {})
        maturity = state.user_inputs.get("cyber_maturity", "Defined")
        user_msg = f"Risk profile: {json.dumps(risks)}\nMaturity: {maturity}\nAdditional: {additional_context}"
        messages = self._build_messages(user_msg)
        raw = self._call_openai(messages, response_format={"type": "json_object"})
        return self._parse_json_response(raw)

class LegalPolicyWriterAgent(BaseAgent):
    def __init__(self, client):
        super().__init__(
            name="Legal & Policy Writer",
            system_prompt=textwrap.dedent("""\
                You are a corporate legal policy drafter. Write formal policy text for each section.
                Output JSON with purpose, scope, definitions, governance_structure, roles_and_responsibilities, policy_statements, operational_controls, monitoring_mechanism, exception_handling, violation_sanctions, kpi_kri.
            """),
            client=client
        )

    def run(self, state: PolicyState, additional_context: str = "") -> Dict[str, Any]:
        prev_outputs = {k: v for k, v in state.agent_outputs.items() if k != self.name}
        user_msg = f"All analysis: {json.dumps(prev_outputs, indent=2)}\nAdditional: {additional_context}"
        messages = self._build_messages(user_msg)
        raw = self._call_openai(messages, response_format={"type": "json_object"})
        return self._parse_json_response(raw)

class AuditReadinessAgent(BaseAgent):
    def __init__(self, client):
        super().__init__(
            name="Audit Readiness",
            system_prompt=textwrap.dedent("""\
                You are an IT auditor. Define audit evidence requirements, KPIs, and monitoring.
                Output JSON.
            """),
            client=client
        )

    def run(self, state: PolicyState, additional_context: str = "") -> Dict[str, Any]:
        policy_draft = state.agent_outputs.get("Legal & Policy Writer", {})
        user_msg = f"Draft policy: {json.dumps(policy_draft)}\nAdditional: {additional_context}"
        messages = self._build_messages(user_msg)
        raw = self._call_openai(messages, response_format={"type": "json_object"})
        return self._parse_json_response(raw)

class ReviewerAgent(BaseAgent):
    def __init__(self, client):
        super().__init__(
            name="Reviewer",
            system_prompt=textwrap.dedent("""\
                You are a quality assurance reviewer. Check for conflicts, ambiguities, overlaps, terminology consistency, feasibility.
                Output JSON with conflicts, ambiguities, overlaps, terminology_issues, feasibility_concerns, overall_quality_score.
            """),
            client=client
        )

    def run(self, state: PolicyState, additional_context: str = "") -> Dict[str, Any]:
        complete_policy = state.agent_outputs.get("Legal & Policy Writer", {})
        user_msg = f"Policy: {json.dumps(complete_policy)}\nAdditional: {additional_context}"
        messages = self._build_messages(user_msg)
        raw = self._call_openai(messages, response_format={"type": "json_object"})
        return self._parse_json_response(raw)

class LegacyPolicyAnalysisAgent(BaseAgent):
    def __init__(self, client):
        super().__init__(
            name="Legacy Policy Analysis",
            system_prompt=textwrap.dedent("""\
                Analyze existing policy document. Identify structure, obsolete clauses, weaknesses, duplicates, governance maturity.
                Output JSON.
            """),
            client=client
        )

    def run(self, state: PolicyState, additional_context: str = "") -> Dict[str, Any]:
        doc_text = state.raw_document[:8000]
        user_msg = f"Existing document (excerpt): {doc_text}\nAdditional: {additional_context}"
        messages = self._build_messages(user_msg)
        raw = self._call_openai(messages, response_format={"type": "json_object"})
        return self._parse_json_response(raw)

class GapAnalysisAgent(BaseAgent):
    def __init__(self, client):
        super().__init__(
            name="Gap Analysis",
            system_prompt=textwrap.dedent("""\
                Perform gap analysis between existing policy and target frameworks.
                Output gap_matrix and summary.
            """),
            client=client
        )

    def run(self, state: PolicyState, additional_context: str = "") -> Dict[str, Any]:
        legacy = state.agent_outputs.get("Legacy Policy Analysis", {})
        frameworks = state.user_inputs.get("main_regulation", "") + ", " + state.user_inputs.get("existing_framework", "")
        user_msg = f"Legacy analysis: {json.dumps(legacy)}\nTarget frameworks: {frameworks}\nAdditional: {additional_context}"
        messages = self._build_messages(user_msg)
        raw = self._call_openai(messages, response_format={"type": "json_object"})
        return self._parse_json_response(raw)

class PolicyModernizationAgent(BaseAgent):
    def __init__(self, client):
        super().__init__(
            name="Policy Modernization",
            system_prompt=textwrap.dedent("""\
                Rewrite and modernize the existing policy. Output modernized_policy (same structure as Legal & Policy Writer), change_log, and modernization_report.
            """),
            client=client
        )

    def run(self, state: PolicyState, additional_context: str = "") -> Dict[str, Any]:
        legacy = state.agent_outputs.get("Legacy Policy Analysis", {})
        gaps = state.agent_outputs.get("Gap Analysis", {})
        user_msg = f"Legacy: {json.dumps(legacy)}\nGap Analysis: {json.dumps(gaps)}\nUser Inputs: {json.dumps(state.user_inputs)}\nAdditional: {additional_context}"
        messages = self._build_messages(user_msg)
        raw = self._call_openai(messages, response_format={"type": "json_object"})
        return self._parse_json_response(raw)

# --------------------------------------------------------------------------------
# Orchestrator
# --------------------------------------------------------------------------------
class PolicyOrchestrator:
    def __init__(self, client: Optional[openai.OpenAI] = None, demo_mode: bool = False):
        self.client = client
        self.demo_mode = demo_mode
        # Always create agents (they may not be used in demo)
        self.agents = {
            "Requirement Analyst": RequirementAnalystAgent(client),
            "Governance & Compliance": GovernanceComplianceAgent(client),
            "Risk Assessment": RiskAssessmentAgent(client),
            "Technical Control": TechnicalControlAgent(client),
            "Legal & Policy Writer": LegalPolicyWriterAgent(client),
            "Audit Readiness": AuditReadinessAgent(client),
            "Reviewer": ReviewerAgent(client),
            "Legacy Policy Analysis": LegacyPolicyAnalysisAgent(client),
            "Gap Analysis": GapAnalysisAgent(client),
            "Policy Modernization": PolicyModernizationAgent(client)
        }

    def run_generate_mode(self, state: PolicyState, progress_callback=None) -> PolicyState:
        execution_order = [
            "Requirement Analyst",
            "Governance & Compliance",
            "Risk Assessment",
            "Technical Control",
            "Legal & Policy Writer",
            "Audit Readiness",
            "Reviewer"
        ]
        if self.demo_mode:
            mock_data = generate_mock_agent_outputs(state)
            for i, agent_name in enumerate(execution_order):
                if progress_callback:
                    progress_callback((i+1)/len(execution_order), f"Simulating {agent_name}...")
                state.agent_outputs[agent_name] = mock_data.get(agent_name, {})
                time.sleep(0.3)
            self._synthesize_policy(state)
        else:
            if not self.client:
                raise RuntimeError("OpenAI client required for live mode.")
            for i, agent_name in enumerate(execution_order):
                if progress_callback:
                    progress_callback((i+1)/len(execution_order), f"Running {agent_name}...")
                agent = self.agents[agent_name]
                try:
                    output = agent.run(state)
                    state.agent_outputs[agent_name] = output
                except Exception as e:
                    state.metadata["errors"].append(f"{agent_name}: {str(e)}")
                    logger.error(f"Agent {agent_name} failed: {e}")
            self._synthesize_policy(state)
        return state

    def run_modernize_mode(self, state: PolicyState, progress_callback=None) -> PolicyState:
        execution_order = [
            "Legacy Policy Analysis",
            "Gap Analysis",
            "Policy Modernization"
        ]
        if self.demo_mode:
            mock_data = generate_mock_agent_outputs(state)
            for i, agent_name in enumerate(execution_order):
                if progress_callback:
                    progress_callback((i+1)/len(execution_order), f"Simulating {agent_name}...")
                state.agent_outputs[agent_name] = mock_data.get(agent_name, {})
                time.sleep(0.3)
            modern_output = state.agent_outputs.get("Policy Modernization", {})
            state.final_policy_json = modern_output.get("modernized_policy", {})
            state.change_log = modern_output.get("change_log", [])
            state.modernization_report = modern_output.get("modernization_report", {})
            state.gap_matrix = state.agent_outputs.get("Gap Analysis", {}).get("gap_matrix", [])
            self._calculate_scores(state)
        else:
            if not self.client:
                raise RuntimeError("OpenAI client required for live mode.")
            for i, agent_name in enumerate(execution_order):
                if progress_callback:
                    progress_callback((i+1)/len(execution_order), f"Running {agent_name}...")
                agent = self.agents[agent_name]
                try:
                    output = agent.run(state)
                    state.agent_outputs[agent_name] = output
                except Exception as e:
                    state.metadata["errors"].append(f"{agent_name}: {str(e)}")
            modern_output = state.agent_outputs.get("Policy Modernization", {})
            state.final_policy_json = modern_output.get("modernized_policy", {})
            state.change_log = modern_output.get("change_log", [])
            state.modernization_report = modern_output.get("modernization_report", {})
            state.gap_matrix = state.agent_outputs.get("Gap Analysis", {}).get("gap_matrix", [])
            self._calculate_scores(state)
        return state

    def _synthesize_policy(self, state: PolicyState):
        legal = state.agent_outputs.get("Legal & Policy Writer", {})
        audit = state.agent_outputs.get("Audit Readiness", {})
        review = state.agent_outputs.get("Reviewer", {})
        compliance = state.agent_outputs.get("Governance & Compliance", {})
        risk = state.agent_outputs.get("Risk Assessment", {})
        tech = state.agent_outputs.get("Technical Control", {})

        state.final_policy_json = {
            "metadata": {
                "policy_name": state.user_inputs.get("policy_name", "Untitled"),
                "organization": state.user_inputs.get("org_type", ""),
                "industry": state.user_inputs.get("industry", ""),
                "version": "1.0",
                "date": datetime.date.today().isoformat(),
                "classification": "Internal"
            },
            "policy_content": legal,
            "audit_readiness": audit,
            "compliance_mapping": compliance.get("compliance_mapping", []),
            "risk_profile": risk,
            "technical_controls": tech,
            "review_feedback": review
        }
        self._calculate_scores(state)

    def _calculate_scores(self, state: PolicyState):
        review = state.agent_outputs.get("Reviewer", {})
        compliance = state.agent_outputs.get("Governance & Compliance", {})
        legacy = state.agent_outputs.get("Legacy Policy Analysis", {})
        modern = state.agent_outputs.get("Policy Modernization", {})

        gaps = compliance.get("regulatory_gaps", [])
        state.scores["compliance"] = max(0.0, 100.0 - len(gaps) * 15)

        mat_level = legacy.get("governance_maturity_estimate", "Defined")
        mat_map = {"Initial": 20, "Repeatable": 40, "Defined": 60, "Managed": 80, "Optimizing": 100}
        state.scores["governance_maturity"] = mat_map.get(mat_level, 60)

        modern_report = modern.get("modernization_report", {})
        improvements = len(modern_report.get("major_improvements", []))
        state.scores["modernization"] = min(100, improvements * 10 + 50)

        quality = review.get("overall_quality_score", 0.8) * 100 if isinstance(review.get("overall_quality_score"), (int, float)) else 70
        state.scores["policy_quality"] = quality

# --------------------------------------------------------------------------------
# Document Generation Utilities
# --------------------------------------------------------------------------------
class DocumentGenerator:
    @staticmethod
    def build_docx(state: PolicyState) -> BytesIO:
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Cover Page
        for _ in range(4):
            doc.add_paragraph()
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(state.final_policy_json.get("metadata", {}).get("policy_name", "POLICY DOCUMENT"))
        run.bold = True
        run.font.size = Pt(26)
        doc.add_paragraph()
        doc.add_paragraph(f"Organization: {state.user_inputs.get('org_type','')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Industry: {state.user_inputs.get('industry','')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Version 1.0 | {datetime.date.today().isoformat()}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()

        # Document Information
        doc.add_heading('Document Information', level=1)
        info = state.final_policy_json.get("metadata", {})
        table = doc.add_table(rows=5, cols=2, style='Light Shading Accent 1')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cells = [
            ("Policy Name", info.get("policy_name", "")),
            ("Version", info.get("version", "1.0")),
            ("Effective Date", info.get("date", "")),
            ("Classification", info.get("classification", "Internal")),
            ("Owner", state.user_inputs.get("org_type", ""))
        ]
        for i, (k, v) in enumerate(cells):
            table.rows[i].cells[0].text = k
            table.rows[i].cells[1].text = v
        doc.add_paragraph()

        # Version Control
        doc.add_heading('Version Control', level=1)
        vtable = doc.add_table(rows=2, cols=4, style='Light Shading Accent 1')
        hdr = vtable.rows[0].cells
        hdr[0].text = 'Version'
        hdr[1].text = 'Date'
        hdr[2].text = 'Author'
        hdr[3].text = 'Changes'
        vtable.rows[1].cells[0].text = '1.0'
        vtable.rows[1].cells[1].text = datetime.date.today().isoformat()
        vtable.rows[1].cells[2].text = 'PolicyAI'
        vtable.rows[1].cells[3].text = 'Initial release'
        doc.add_paragraph()

        # Approval
        doc.add_heading('Approval', level=1)
        doc.add_paragraph('This document has been approved by:')
        doc.add_paragraph('[Name] - [Role]')
        doc.add_paragraph('Date: _______________')
        doc.add_page_break()

        # Policy Content
        content = state.final_policy_json.get("policy_content", {})
        sections = [
            ("Purpose", content.get("purpose", "")),
            ("Scope", content.get("scope", "")),
            ("Definitions", "\n".join([f"{k}: {v}" for k,v in content.get("definitions", {}).items()])),
            ("Governance Structure", content.get("governance_structure", "")),
            ("Roles & Responsibilities", content.get("roles_and_responsibilities", "")),
            ("Policy Statements", "\n".join(content.get("policy_statements", []))),
            ("Operational Controls", content.get("operational_controls", "")),
            ("Monitoring Mechanism", content.get("monitoring_mechanism", "")),
            ("Exception Handling", content.get("exception_handling", "")),
            ("Violation & Sanctions", content.get("violation_sanctions", "")),
            ("KPI/KRI", "\n".join(content.get("kpi_kri", [])))
        ]
        for heading, text in sections:
            if text:
                doc.add_heading(heading, level=1)
                doc.add_paragraph(text)
        doc.add_page_break()

        # Technical Controls
        tech = state.final_policy_json.get("technical_controls", {})
        if tech:
            doc.add_heading('Technical Controls', level=1)
            for k, v in tech.items():
                doc.add_heading(k.replace('_', ' ').title(), level=2)
                doc.add_paragraph(str(v) if isinstance(v, str) else "\n".join(v))
        doc.add_page_break()

        # Compliance Mapping
        comp = state.final_policy_json.get("compliance_mapping", [])
        if comp:
            doc.add_heading('Compliance Mapping', level=1)
            ctable = doc.add_table(rows=1, cols=4, style='Light Shading Accent 1')
            hdr = ctable.rows[0].cells
            hdr[0].text = 'Framework'
            hdr[1].text = 'Clause'
            hdr[2].text = 'Control ID'
            hdr[3].text = 'Status'
            for item in comp:
                row = ctable.add_row()
                row.cells[0].text = item.get('framework', '')
                row.cells[1].text = item.get('clause', '')
                row.cells[2].text = item.get('control_id', '')
                row.cells[3].text = item.get('status', '')

        # Audit Evidence
        audit = state.final_policy_json.get("audit_readiness", {}).get("evidence_requirements", [])
        if audit:
            doc.add_heading('Audit Evidence Requirements', level=1)
            atable = doc.add_table(rows=1, cols=3, style='Light Shading Accent 1')
            hdr = atable.rows[0].cells
            hdr[0].text = 'Control Area'
            hdr[1].text = 'Evidence Type'
            hdr[2].text = 'Frequency'
            for ev in audit:
                row = atable.add_row()
                row.cells[0].text = ev.get('control_area', '')
                row.cells[1].text = ev.get('evidence_type', '')
                row.cells[2].text = ev.get('frequency', '')

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    @staticmethod
    def build_pdf(state: PolicyState) -> BytesIO:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(0, 10, state.final_policy_json.get("metadata", {}).get("policy_name", "Policy"), ln=True, align='C')
        pdf.ln(10)
        pdf.set_font("Arial", size=10)
        for heading, text in state.final_policy_json.get("policy_content", {}).items():
            if text and isinstance(text, str):
                pdf.set_font("Arial", 'B', 12)
                pdf.cell(0, 10, heading.replace('_', ' ').title(), ln=True)
                pdf.set_font("Arial", size=10)
                pdf.multi_cell(0, 5, text)
                pdf.ln(5)
        buffer = BytesIO()
        pdf.output(buffer)
        buffer.seek(0)
        return buffer

    @staticmethod
    def build_markdown(state: PolicyState) -> str:
        md = f"# {state.final_policy_json.get('metadata', {}).get('policy_name', 'Policy')}\n\n"
        content = state.final_policy_json.get("policy_content", {})
        for heading, text in content.items():
            if text:
                md += f"## {heading.replace('_', ' ').title()}\n\n{text}\n\n"
        return md

# --------------------------------------------------------------------------------
# UI Components
# --------------------------------------------------------------------------------
def render_sidebar(state: PolicyState):
    with st.sidebar:
        st.image("https://via.placeholder.com/150x50?text=PolicyAI", width=150)
        st.markdown("### Configuration")
        api_key = st.text_input("OpenAI API Key", type="password", help="Enter your OpenAI API key for live AI mode")
        if api_key:
            st.session_state["openai_api_key"] = api_key
            st.session_state.demo_mode = False
        else:
            st.warning("⚠️ No API key. You can run in Demo Mode.")
            demo = st.checkbox("Enable Demo Mode (simulated AI)", value=True)
            st.session_state.demo_mode = demo
            if demo:
                st.success("Demo Mode active – all outputs are simulated.")

        st.divider()
        st.markdown("### Application Mode")
        mode = st.radio("Select Mode", ["Generate New Policy", "Modernize Existing Policy"],
                         captions=["Create policy from scratch", "Update legacy document"])
        state.mode = "generate" if mode.startswith("Generate") else "modernize"

        st.divider()
        st.markdown("### Policy Configuration")
        with st.form("policy_form"):
            policy_name = st.text_input("Policy Name *", placeholder="e.g., Information Security Policy")
            org_type = st.selectbox("Organization Type *", Config.INDUSTRIES)
            industry = st.selectbox("Industry *", Config.INDUSTRIES)
            main_reg = st.multiselect("Main Regulation/Framework *", Config.FRAMEWORKS, default=["ISO 27001"])
            cyber_maturity = st.select_slider("Cybersecurity Maturity", Config.MATURITY_LEVELS, value="Defined")
            org_scale = st.selectbox("Organization Scale", Config.ORG_SCALES)
            it_complexity = st.select_slider("IT Complexity", ["Low", "Medium", "High", "Very High"], value="Medium")
            business_model = st.text_input("Business Model", placeholder="e.g., B2B SaaS")
            existing_framework = st.multiselect("Existing Adopted Frameworks", Config.FRAMEWORKS)
            policy_objective = st.text_area("Policy Objective", placeholder="Primary goal of this policy")
            additional_notes = st.text_area("Additional Notes", placeholder="Any specific requirements or context")

            uploaded_file = None
            if state.mode == "modernize":
                uploaded_file = st.file_uploader("Upload Existing Policy Document",
                                                 type=Config.SUPPORTED_INPUT_FORMATS,
                                                 accept_multiple_files=False)

            submitted = st.form_submit_button("Generate Policy", type="primary", use_container_width=True)

        if submitted:
            if not policy_name or not org_type:
                st.error("Policy Name and Organization Type are required.")
                return False
            state.user_inputs = {
                "policy_name": policy_name,
                "org_type": org_type,
                "industry": industry,
                "main_regulation": ", ".join(main_reg),
                "cyber_maturity": cyber_maturity,
                "org_scale": org_scale,
                "it_complexity": it_complexity,
                "business_model": business_model,
                "existing_framework": ", ".join(existing_framework),
                "policy_objective": policy_objective,
                "additional_notes": additional_notes
            }
            if uploaded_file is not None:
                file_bytes = uploaded_file.read()
                file_type = uploaded_file.name.split('.')[-1].lower()
                state.user_inputs["uploaded_file_name"] = uploaded_file.name
                try:
                    state.raw_document = DocumentParser.parse_document(file_bytes, file_type)
                    chunks = SemanticChunker.chunk_by_headings(state.raw_document)
                    if not chunks:
                        chunks = [{"heading": "Full Document", "content": state.raw_document}]
                    state.chunks = chunks
                    st.success(f"Document parsed: {uploaded_file.name}")
                except Exception as e:
                    st.error(f"Document parsing failed: {e}")
                    return False
            return True
        return False

def main_content(state: PolicyState):
    st.title("🔒 PolicyAI Enterprise")
    st.caption("Agentic AI Governance & Policy Engineering Platform")

    # Determine mode
    demo_mode = st.session_state.get("demo_mode", False)
    client = get_openai_client()

    if not client and not demo_mode:
        st.error("Please enter a valid OpenAI API key or enable Demo Mode in the sidebar.")
        st.stop()

    if not state.user_inputs:
        st.info("👈 Configure policy parameters in the sidebar and click 'Generate Policy' to start.")
        return

    orchestrator = PolicyOrchestrator(client=client if client else None, demo_mode=demo_mode)

    progress_bar = st.progress(0, text="Initializing...")

    def progress_update(percent, message):
        progress_bar.progress(percent, text=message)

    if state.mode == "generate":
        with st.spinner("Agents working on policy generation..."):
            try:
                state = orchestrator.run_generate_mode(state, progress_callback=progress_update)
            except Exception as e:
                st.error(f"Workflow failed: {e}")
                logger.exception("Orchestrator error")
                st.stop()
    else:
        with st.spinner("Analyzing and modernizing existing policy..."):
            try:
                state = orchestrator.run_modernize_mode(state, progress_callback=progress_update)
            except Exception as e:
                st.error(f"Modernization workflow failed: {e}")
                logger.exception("Orchestrator error")
                st.stop()

    progress_bar.progress(1.0, text="Complete!")
    time.sleep(0.5)
    progress_bar.empty()

    st.success("✅ Policy generation complete!")

    # Scores
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Compliance Score", f"{state.scores['compliance']:.0f}%")
    with col2:
        st.metric("Modernization Score", f"{state.scores['modernization']:.0f}%")
    with col3:
        st.metric("Governance Maturity", f"{state.scores['governance_maturity']:.0f}%")
    with col4:
        st.metric("Policy Quality", f"{state.scores['policy_quality']:.0f}%")

    st.divider()

    # Agent Log
    with st.expander("🔍 Agent Execution Log", expanded=False):
        for agent_name, output in state.agent_outputs.items():
            st.subheader(agent_name)
            st.json(output, expanded=False)

    # Tabs
    tab1, tab2, tab3 = st.tabs(["📄 Generated Policy", "📊 Gap Analysis & Reports", "🔄 Change Log"])

    with tab1:
        st.subheader("Final Policy Document")
        st.json(state.final_policy_json, expanded=False)

        col_d1, col_d2, col_d3, col_d4 = st.columns(4)
        with col_d1:
            docx_buffer = DocumentGenerator.build_docx(state)
            st.download_button("⬇ Download DOCX", data=docx_buffer, file_name="policy.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        with col_d2:
            pdf_buffer = DocumentGenerator.build_pdf(state)
            st.download_button("⬇ Download PDF", data=pdf_buffer, file_name="policy.pdf",
                               mime="application/pdf")
        with col_d3:
            md_content = DocumentGenerator.build_markdown(state)
            st.download_button("⬇ Download Markdown", data=md_content, file_name="policy.md",
                               mime="text/markdown")
        with col_d4:
            json_str = json.dumps(state.final_policy_json, indent=2)
            st.download_button("⬇ Download JSON", data=json_str, file_name="policy.json",
                               mime="application/json")

    with tab2:
        if state.mode == "modernize":
            st.subheader("Modernization Report")
            report = state.modernization_report
            if report:
                st.markdown(f"**Executive Summary:** {report.get('executive_summary', 'N/A')}")
                st.markdown("**Major Improvements:**")
                for imp in report.get("major_improvements", []):
                    st.markdown(f"- {imp}")
                col_a, col_b = st.columns(2)
                with col_a:
                    st.metric("Governance Improvement", report.get("governance_improvement", "N/A"))
                    st.metric("Compliance Improvement", report.get("compliance_improvement", "N/A"))
                with col_b:
                    st.metric("Security Improvement", report.get("security_improvement", "N/A"))
                    st.metric("Audit Readiness", report.get("audit_readiness_improvement", "N/A"))
                st.markdown("**Recommended Next Actions:**")
                for action in report.get("next_actions", []):
                    st.markdown(f"- {action}")

        st.subheader("Gap Analysis Matrix")
        gaps = state.gap_matrix
        if gaps:
            df = pd.DataFrame(gaps)
            st.dataframe(df, use_container_width=True)
        else:
            st.info("No gap matrix available for this mode.")

    with tab3:
        if state.change_log:
            st.subheader("Policy Change Log")
            df = pd.DataFrame(state.change_log)
            st.dataframe(df, use_container_width=True)
            if st.checkbox("Show side-by-side diff"):
                for change in state.change_log:
                    col_old, col_new = st.columns(2)
                    with col_old:
                        st.markdown("**Old Clause**")
                        st.info(change.get("old_clause", ""))
                    with col_new:
                        st.markdown("**New Clause**")
                        st.success(change.get("new_clause", ""))
                    st.caption(f"Reason: {change.get('reason','')} | Impact: {change.get('compliance_impact','')}")
                    st.divider()
        else:
            st.info("No change log generated (only for modernization mode).")

# --------------------------------------------------------------------------------
# Main App
# --------------------------------------------------------------------------------
def main():
    st.set_page_config(
        page_title="PolicyAI Enterprise",
        page_icon="🔒",
        layout="wide",
        initial_sidebar_state="expanded",
        menu_items={
            'Get Help': None,
            'Report a bug': None,
            'About': "PolicyAI Enterprise v2.0 - AI Governance Platform"
        }
    )

    # Initialize session state
    if "policy_state" not in st.session_state:
        st.session_state.policy_state = PolicyState()
    if "demo_mode" not in st.session_state:
        st.session_state.demo_mode = True  # default demo

    state = st.session_state.policy_state

    new_gen = render_sidebar(state)
    if new_gen:
        # Reset state for fresh generation
        st.session_state.policy_state = PolicyState()
        state = st.session_state.policy_state
        # Re-populate user inputs from sidebar (done in render_sidebar)

    main_content(state)

if __name__ == "__main__":
    main()
